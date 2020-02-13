#!/usr/bin/env python3

########################################################################
# Python 3                                               Quentin Petit #
# February 2020                                <petit.quent@gmail.com> #
#                                                                      #
#                             rst2docx.py                              #
#                                                                      #
# Current version: 1.0.0                                               #
# Status: Work in progress                                             #
#                                                                      #
# This script purpose it to format reStructuredText notes to MS docx   #
# file.                                                                #
#                                                                      #
# Version history:                                                     #
# +----------+---------+---------------------------------------------+ #
# |   Date   | Version | Comment                                     | #
# +==========+=========+=============================================+ #
# | 20200207 | 0.1.0   | First development                           | #
# +----------+---------+---------------------------------------------+ #
# | 20200213 | 1.0.0   | Fist stable version                         | #
# +----------+---------+---------------------------------------------+ #
#                                                                      #
# The prerequisites to use this script are:                            #
#                                                                      #
# o  The Python package python-docx should be installed                #
#                                                                      #
########################################################################

#                                                                      #
#                              LIBRAIRIES                              #
#                                                                      #

import argparse
import configparser
import os
import re
import shutil
import sys
from docx import Document
from docx.shared import Inches

#~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~#
#                                                                      #
#                               VARIABLES                              #
#                                                                      #

# Files and directories.
dir_current = os.path.abspath(os.path.dirname(sys.argv[0]))
dir_root = dir_current.rsplit('/', 1)[0]
dir_ini = dir_root + "/ini/"
dir_tmp = dir_root + "/tmp/"
docx_template = dir_ini + "tpl_rst2docx.docx"
file_tmpprgrph = dir_tmp + "tmp_paragraph"
file_tmpprgrphold = dir_tmp + "tmp_paragraph_previous"
file_ini = dir_ini + "rst2docx.ini"

# Lists and symbols.
nonalphanum = ["=","-","`",":","'","\"","~","^","_","*","+","#","<",">"]
bullet_symbol = ["*","-","+"]
admonition_drctves = ["ATTENTION","CAUTION","DANGER","ERROR","HINT","IMPORTANT","NOTE","TIP","WARNING","ADMONITION"]
attention_drctves = ["ATTENTION","CAUTION","WARNING"]
danger_drctves = ["DANGER","ERROR"]
hint_drctves = ["HINT","IMPORTANT","TIP"]
note_drctves = ["NOTE","ADMONITION"]

# Declaration of variables
separator = " "
header_symbol = []

#~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~#
#                                                                      #
#                               FUNCTIONS                              #
#                                                                      #

# Function to print error when title not found.
def title_notfound():
    print("error: not title found in " + os.path.basename(src))
    print("see https://docutils.readthedocs.io/en/sphinx-docs/user/rst/quickstart.html#document-title-subtitle for more details")
    sys.exit(1)

# Function to apply text inside paragraph
def txt_style(p, split_prgrph):
    word_tag = "none"
    for word in split_prgrph:
        if word == split_prgrph[-1] and word.endswith("::"):
            word = word[:-1]
        if word_tag == "none":
            if word.startswith("*"):
                if word.startswith("**"):
                    if word.endswith("**"):
                        p.add_run(word[:-2][2:] + " ").bold = True
                    else:
                        word_tag = "bold"
                        p.add_run(word[2:] + " ").bold = True
                else:
                    if word.endswith("*"):
                        p.add_run(word[:-1][1:] + " ").italic = True
                    else:
                        word_tag = "italics"
                        p.add_run(word[1:] + " ").italic = True
            elif word.startswith("``"):
                if word.endswith("``"):
                    p.add_run(word[:-2][2:]).style = 'InlineCode'
                    p.add_run(" ").style = None
                else:
                    word_tag = "code"
                    p.add_run(word[2:] + " ").style = 'InlineCode'
            else:
                p.add_run(word + " ")
        elif word_tag == "bold":
            if word.endswith("**"):
                word_tag = "none"
                p.add_run(word[:-2] + " ").bold = True
            else:
                p.add_run(word + " ").bold = True
        elif word_tag == "italics":
            if word.endswith("*"):
                word_tag = "none"
                p.add_run(word[:-1] + " ").italic = True
            else:
                p.add_run(word + " ").italic = True
        elif word_tag == "code":
            if word.endswith("``"):
                word_tag = "none"
                p.add_run(word[:-2]).style = 'InlineCode'
                p.add_run(" ").style = None
            else:
                p.add_run(word + " ").style = 'InlineCode'

# Function to write a paragraph from RST.
def write_prgrph():
    file_tmp.close()
    with open(file_tmpprgrph, "r") as tmp_prgrph:
        prgrph = tmp_prgrph.read()
        with open(file_tmpprgrphold, "r") as tmp_prgrphold:
            prgrph_old = tmp_prgrphold.read()
            if prgrph != prgrph_old:
                # Headers.
                lines = prgrph.split('\n')
                try:
                    lines.remove('')
                except:
                    pass
                if len(lines) == 2 and re.match('^[=\-`:\'"~^_*+#<>]+$', lines[1]):
                    if lines[1][0] not in header_symbol:
                        header_symbol.append(lines[1][0])
                    header_lvl = header_symbol.index(lines[1][0]) + 1
                    document.add_paragraph(lines[0], 'Header_' + str(header_lvl))
                # Code blocks.
                elif prgrph_old.rstrip('\n')[-2:] == "::":
                    for line in lines:
                        document.add_paragraph(line[2:], 'Code')
                # Python code blocks.
                elif prgrph.startswith(">>> "):
                    document.add_paragraph(prgrph.rstrip(), 'Code')
                # Bulleted lists.
                elif lines[0].strip()[0] in bullet_symbol and lines[0].strip()[1] == " ":
                    lead_space = len(lines[0]) - len(lines[0].strip())
                    bullet_lvl = int(lead_space / 2) + 1
                    document.add_paragraph(separator.join(lines).strip()[2:], 'Bullet_' + str(bullet_lvl))
                # Enumerated lists.
                elif re.match('^[(\d\w]+[).]$', lines[0].strip().split()[0]):
                    lead_space = len(lines[0]) - len(lines[0].strip())
                    number_lvl = int(lead_space / 3) + 1
                    document.add_paragraph(separator.join(separator.join([x.strip() for x in lines]).split()[1:]), 'Number_' + str(number_lvl))
                # Admonitions.
                elif lines[0].split()[0] == ".." and lines[0].split()[1][:-2] in admonition_drctves:
                    if lines[0].split()[1][:-2] in attention_drctves:
                        p = document.add_paragraph("", 'ATTENTION')
                        p.add_run(lines[0].split()[1][:-2] + "\n").bold = True
                        split_prgrph = prgrph.split()[2:]
                        txt_style(p, split_prgrph)
                    elif lines[0].split()[1][:-2] in danger_drctves:
                        p = document.add_paragraph("", 'DANGER')
                        p.add_run(lines[0].split()[1][:-2] + "\n").bold = True
                        split_prgrph = prgrph.split()[2:]
                        txt_style(p, split_prgrph)
                    elif lines[0].split()[1][:-2] in hint_drctves:
                        p = document.add_paragraph("", 'HINT')
                        p.add_run(lines[0].split()[1][:-2] + "\n").bold = True
                        split_prgrph = prgrph.split()[2:]
                        txt_style(p, split_prgrph)
                    elif lines[0].split()[1][:-2] in note_drctves:
                        p = document.add_paragraph("", 'NOTE')
                        p.add_run(lines[0].split()[1][:-2] + "\n").bold = True
                        split_prgrph = prgrph.split()[2:]
                        txt_style(p, split_prgrph)
                # Tables.
                elif prgrph.startswith("+-") and re.match('^[=\-+]+$', lines[0]) and re.match('^[=\-+]+$', lines[-1]):
                    # Reset variables for tables.
                    tab_pipe_nb = []
                    tab_vals = []
                    tab_vals_len = []
                    merge_tag = 0
                    tab_row_nb = 0
                    row_ix = 0
                    # Get number of columns in table.
                    for line in lines:
                        tab_pipe_nb.append(line.count('|'))
                    tab_col_nb = max(tab_pipe_nb) - 1
                    # Get nb of rows and values in table.
                    for line in lines:
                        if "+-" not in line and "-+" not in line and "+=" not in line and "=+" not in line:
                            tab_row_nb += 1
                            row_vals = line.split("|")
                            row_vals.pop(0)
                            row_vals.pop(-1)
                            row_vals_len = [len(x) for x in row_vals]
                            row_vals = [x.strip(' ') for x in row_vals]
                            tab_vals.append(row_vals)
                            tab_vals_len.append(row_vals_len)
                    def_cell_len = min(tab_vals_len)
                    # Init table.
                    table = document.add_table(tab_row_nb, tab_col_nb, 'Table_Light')
                    # Fill table.
                    for row_vals in tab_vals:
                        col_ix = 0
                        if len(row_vals) == tab_col_nb:
                            for value in row_vals:
                                if value != '':
                                    table.cell(row_ix, col_ix).text = value
                                else:
                                    table.cell(row_ix - 1, col_ix).merge(table.cell(row_ix, col_ix))
                                col_ix += 1
                        elif len(row_vals) == 1: 
                            table.cell(row_ix, 0).text = row_vals[0]
                            col_ix = 1
                            while col_ix != tab_col_nb:
                                table.cell(row_ix, col_ix - 1).merge(table.cell(row_ix, col_ix))
                                col_ix += 1
                        else:
                            col_ix_src = col_ix
                            while col_ix != tab_col_nb:
                                if merge_tag == 1:
                                    table.cell(row_ix, col_ix - 1).merge(table.cell(row_ix, col_ix))
                                    len_merged_cell = tab_vals_len[row_ix][col_ix - 1]
                                    current_merged_cell_len = def_cell_len[col_ix] + def_cell_len[col_ix - 1] + 1
                                    col_ix += 1
                                    while len_merged_cell != current_merged_cell_len:
                                        table.cell(row_ix, col_ix - 1).merge(table.cell(row_ix, col_ix))
                                        current_merged_cell_len = current_merged_cell_len + def_cell_len[col_ix] + 1
                                        col_ix += 1
                                    merge_tag = 0
                                elif def_cell_len[col_ix] == tab_vals_len[row_ix][col_ix_src]:
                                    table.cell(row_ix, col_ix).text = row_vals[col_ix_src]
                                    merge_tag = 0
                                    col_ix += 1
                                    col_ix_src += 1
                                elif def_cell_len[col_ix] != tab_vals_len[row_ix][col_ix_src]:
                                    table.cell(row_ix, col_ix).text = row_vals[col_ix_src]
                                    merge_tag = 1
                                    col_ix += 1
                                    col_ix_src += 1
                        row_ix += 1
                # Sources.
                elif lines[0].startswith(".. ["):
                    table = document.add_table(1, 2)
                    table.cell(0, 0).text = prgrph.split()[1]
                    table.cell(0, 0).width = Inches(0.8)
                    table.cell(0, 1).text = separator.join(prgrph.split()[2:])
                    table.cell(0, 1).width = Inches(5.3)
                # Markup for code.
                elif lines[0] == "::" and len(lines) == 1:
                    pass
                # Normal text.
                else:
                    p = document.add_paragraph("", 'Normal')
                    split_prgrph = prgrph.split()
                    txt_style(p, split_prgrph)
    shutil.copyfile(file_tmpprgrph, file_tmpprgrphold)

#~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~#
#                                                                      #
#                               BEGINNING                              #
#                                                                      #

# Retrieve arguments.
parser = argparse.ArgumentParser(description='Format documentation')
parser.add_argument('-s', action="store", 
                    dest="src", help="source file")
args = parser.parse_args()

# Check if all arguments are set.
if len(sys.argv) > 1 and args.src == "":
    print("error: one or more argument is missing")
    parser.print_help()
    sys.exit(1)
# If no argument is set, interactive session.
elif len(sys.argv) <= 1:
    src = input("Source document: ")
else:
    src = args.src

# Print error and quit if file doesn't exist.
if not os.path.isfile(src):
    print("error: file " + src + " doesn't exist, bye")
    sys.exit(1)

# Print error and quit if ini file doesn't exist.
if not os.path.isfile(file_ini):
    print("error: ini file " + file_ini + " doesn't exist, bye")
    sys.exit(1)

# Open config file.
config = configparser.ConfigParser()
config.read(file_ini)

# Create directories if not existing.
if not os.path.isdir(dir_tmp):
    os.makedirs(dir_tmp)

# Set output filename.
docx_output = dir_tmp + os.path.basename(src).replace('.rst', '.docx')

# Open docx template.
document = Document(docx_template)

# Retrieve and write title and subtitle.
with open(src, "r") as src_file:
    file_tmp = open(file_tmpprgrph, 'w')
    for line in src_file:
        if line == "\n":
            file_tmp.close()
            count = 0
            with open(file_tmpprgrph, 'r') as tmp_prgrph:
                for line in tmp_prgrph:
                    count += 1
            with open(file_tmpprgrph, 'r') as tmp_prgrph:
                if count == 3:
                    count = 0
                    for line in tmp_prgrph:
                        count += 1
                        if (count % 2) != 0 and line[0] not in nonalphanum:
                            title_notfound()
                        else:
                            title = re.sub(r"^\s+", "", line.rstrip())
                            subtitle = ""
                            document.add_paragraph(title, 'Cover_Title')
                elif count == 6:
                    count = 0
                    for line in tmp_prgrph:
                        count += 1
                        if str(count) in ["1","3","4","6"] and line[0] not in nonalphanum:
                            title_notfound()
                        elif str(count) == "2":
                            title = re.sub(r"^\s+", "", line.rstrip())
                            document.add_paragraph(title, 'Cover_Title')
                        elif count == 5:
                            subtitle = re.sub(r"^\s+", "", line.rstrip())
                            document.add_paragraph(subtitle, 'Cover_Subtitle')
                else:
                    title_notfound()
            shutil.copyfile(file_tmpprgrph, file_tmpprgrphold)
            break
        else:
            file_tmp.write(line)

# Insert page break.
document.add_page_break()

# Insert disclaimer.
document.add_paragraph("Disclaimer", 'Header_A')
document.add_paragraph(re.sub('\\n(?!\\n)' , ' ', config['DEFAULT']['Disclaimer']).replace('\n ', '\n\n'), 'Normal')

# Insert writing conventions.
document.add_paragraph("Conventions", 'Header_A')
document.add_paragraph("The following conventions are used in this document:", 'Normal')
document.add_paragraph("Code blocks, commands or file content:", 'Bullet_1')
document.add_paragraph("echo \"Hello Wold\"", 'Code')
p = document.add_paragraph("The following text is a variable that should be replaced by it correct value: ", 'Bullet_1')
p.add_run('${VERSION}').style = 'InlineCode'
p.add_run(" ").style = None

# Insert content.
with open(src, "r") as src_file:
    file_tmp = open(file_tmpprgrph, 'w')
    for line in src_file:
        if line == "\n":
            write_prgrph()
            file_tmp = open(file_tmpprgrph, 'w')
        else:
            file_tmp.write(line)
write_prgrph()

# Insert authors details.
document.add_paragraph("Authors", 'Header_A')
document.add_paragraph(config['User']['Full_name'] + " (editor)", 'No_Spacing')
document.add_paragraph(config['User']['Corp'], 'No_Spacing')
document.add_paragraph(config['User']['Corp_Address'], 'No_Spacing')
document.add_paragraph("EMail: " + config['User']['Email'], 'No_Spacing')

# Save document.
document.save(docx_output)

# Cleanup
os.remove(file_tmpprgrph)
os.remove(file_tmpprgrphold)

#                                                                      #
#                                  END                                 #
#                                                                      #
#~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~#
