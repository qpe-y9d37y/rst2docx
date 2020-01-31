#!/usr/bin/env python3

########################################################################
# Python 3                                               Quentin Petit #
#                                              <petit.quent@gmail.com> #
#                                                         January 2020 #
#                                                                      #
#                                                                      #
#                                                                      #
#                             nir2docx.py                              #
#                                                                      #
#                                                                      #
#                                                                      #
# Current version: 0.1                                                 #
# Status: Work in progress                                             #
#                                                                      #
# This script purpose it to format NIR notes to MS docx file.          #
#                                                                      #
#                                                                      #
#                                                                      #
# Version history:                                                     #
# +----------+------------+------+-----------------------------------+ #
# |   Date   |   Author   | Vers | Comment                           | #
# +----------+------------+------+-----------------------------------+ #
# | 20200130 | Quentin P. | 0.1  | Starting development              | #
# +----------+------------+------+-----------------------------------+ #
#                                                                      #
########################################################################

#                                                                      #
#                              LIBRAIRIES                              #
#                                                                      #

import argparse
import os
import re
import sys
from docx import Document

#~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~#
#                                                                      #
#                               VARIABLES                              #
#                                                                      #

# Files and directories
dir_current = os.path.abspath(os.path.dirname(sys.argv[0]))
dir_root = dir_current.rsplit('/', 1)[0]
dir_ini = dir_root + "/ini/"
dir_tmp = dir_root + "/tmp/"
docx_template = dir_ini + "tpl_nir2docx.docx"

# Import all variables from file
sys.path.append(dir_ini)
from var_nirutils import *

#~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~#
#                                                                      #
#                               FUNCTIONS                              #
#                                                                      #

def write_prgrph():
    file_tmp.close()
    with open (dir_tmp + "tmp_paragraph", "r") as tmp_prgrph:
        prgrph = tmp_prgrph.read()
        if prgrph.startswith("."):
            if prgrph.startswith(".."):
                prgrph = prgrph.replace('..  ', '').rstrip()
                document.add_paragraph(prgrph, 'Header_2')
            else:
                prgrph = prgrph.replace('.  ', '').rstrip()
                document.add_paragraph(prgrph, 'Header_1')
        elif prgrph.startswith("   "):
            if prgrph.startswith("        "):
                prgrph = prgrph.replace('        ', '').rstrip()
                document.add_paragraph(prgrph, 'Code')
            elif prgrph.startswith("   o  "):
                prgrph = prgrph.replace('   o  ', '').replace('      ', '').replace("\n", " ").rstrip()
                document.add_paragraph(prgrph, 'Bullet_1')
            else:
                prgrph = prgrph.replace('   ', '').replace("\n", " ").rstrip()
                document.add_paragraph(prgrph, 'Normal')
    os.remove(dir_tmp + "tmp_paragraph")

#~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~#
#                                                                      #
#                               BEGINNING                              #
#                                                                      #

# Retrieve arguments.
parser = argparse.ArgumentParser(description='Format documentation')
parser.add_argument('-s', action="store", 
                    dest="src", help="source file")
parser.add_argument('-T', action="store", dest="title", help="title")
parser.add_argument('-c', action="store", dest="cust", help="customer")
args = parser.parse_args()

# Check if all arguments are set.
if len(sys.argv) > 1 and \
   None in (args.src, args.title, args.cust):
    print("error: one or more argument is missing")
    parser.print_help()
    sys.exit(1)
# If no argument is set, interactive session.
elif len(sys.argv) <= 1:
    src = input("Source document: ")
    title = input("Document title: ")
    cust = input("Customer: ")
else:
    src = args.src
    title = args.title
    cust = args.cust

# Print error and quit if file doesn't exist.
if not os.path.isfile(src):
    print("error: file " + src + " doesn't exist, bye")
    sys.exit(1)

# Create directories if not existing.
if not os.path.isdir(dir_tmp):
    os.makedirs(dir_tmp)

# Set output filename
docx_output = dir_tmp + os.path.basename(src).replace('.nir', '.docx')

# Open docx template
document = Document(docx_template)

# Write cover
document.add_paragraph(title, 'Cover_Title')
document.add_paragraph(cust, 'Cover_Subtitle')
# Insert page break
document.add_page_break()

# Insert disclaimer
document.add_paragraph("Disclaimer:", 'Cover_Subtitle')
disclaimer = re.sub('\\n(?!\\n)' , ' ',
                    disclaimer.rstrip()).replace("\n", "\n\n")
########################################################################
# TO CORRECT: Leading space at the beginning of each sentence.         #
########################################################################
document.add_paragraph(disclaimer, 'Normal')

# Insert content
with open(src, "r") as src_file:
    file_tmp = open(dir_tmp + "tmp_paragraph", 'w')
    for line in src_file:
        if line == "\n":
            write_prgrph()
            file_tmp = open(dir_tmp + "tmp_paragraph", 'w')
        else:
            file_tmp.write(line)
write_prgrph()

# Insert authors details
document.add_paragraph("Authors", 'Header_A')
document.add_paragraph(fname + " " + lname + " (editor)", 'No_Spacing')
if corp != "":
    document.add_paragraph(corp, 'No_Spacing')
if corpad != "":
    document.add_paragraph(corpad, 'No_Spacing')
document.add_paragraph("EMail: " + email, 'No_Spacing')

# Save document
document.save(docx_output)

#                                                                      #
#                                  END                                 #
#                                                                      #
#~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~#
